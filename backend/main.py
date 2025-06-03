from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import json
import os
import datetime
from typing import List, Dict, Union
import shutil
from fastapi.responses import FileResponse
import docx
from docx import Document
import chardet
import textract
import mammoth
import PyPDF2
import io
import requests
from screen import (
    extract_text_from_resume,
    clean_resume_text,
    analyze_resume,
    test_ollama_connection
)

# Define the path to the users JSON file
USERS_FILE = os.path.join(os.path.dirname(__file__), 'users.json')

# Initialize FastAPI app
app = FastAPI()

# Configure CORS
# In a production environment, you should restrict origins to your frontend URL
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "http://localhost:3002"], # Allow requests from your frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models for request data
class SeekerSignupData(BaseModel):
    role: str = "seeker"
    name: str # Full Name
    email: str
    password: str
    re_password: str
    phone: str
    job_title: str
    experience: str
    skills: List[str]
    location: str
    resume_url: str

class EmployerSignupData(BaseModel):
    role: str = "employer"
    name: str # Company Name
    email: str
    password: str
    re_password: str
    phone: str
    industry: str
    company_size: str
    company_website: str
    company_location: str
    company_description: str
    contact_person: str

class SignupRequest(BaseModel):
    role: str
    email: str
    password: str
    re_password: str
    name: str
    phone: str
    job_title: Union[str, None] = None
    experience: Union[str, None] = None
    skills: Union[List[str], None] = None
    location: Union[str, None] = None
    resume_url: Union[str, None] = None
    industry: Union[str, None] = None
    company_size: Union[str, None] = None
    company_website: Union[str, None] = None
    company_location: Union[str, None] = None
    company_description: Union[str, None] = None
    contact_person: Union[str, None] = None

# Pydantic model for login request data
class LoginRequest(BaseModel):
    email: str
    password: str

# Configure folders for resume screening
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'Original_Resumes')
PARSED_RESUMES_FOLDER = os.path.join(BASE_DIR, 'parsed_resumes')
JOB_DESCRIPTION_FOLDER = os.path.join(BASE_DIR, 'Job_Description')

# Create folders if they don't exist
for folder in [UPLOAD_FOLDER, PARSED_RESUMES_FOLDER, JOB_DESCRIPTION_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

def cleanup_parsed_resumes():
    if os.path.exists(PARSED_RESUMES_FOLDER):
        shutil.rmtree(PARSED_RESUMES_FOLDER)
    os.makedirs(PARSED_RESUMES_FOLDER)

@app.middleware("http")
async def cleanup_middleware(request, call_next):
    cleanup_parsed_resumes()
    response = await call_next(request)
    return response

# Configure allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'rtf', 'odt', 'jpg', 'jpeg', 'png'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_word(file_path):
    """Extract text from Word documents (.doc and .docx)."""
    try:
        if file_path.endswith('.docx'):
            try:
                doc = Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
            except Exception as e:
                print(f"Error reading .docx file: {str(e)}")
                return None
        elif file_path.endswith('.doc'):
            try:
                # First try using python-docx
                doc = Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
            except Exception as e:
                print(f"Error reading .doc file with python-docx: {str(e)}")
                try:
                    # Fallback to mammoth
                    with open(file_path, 'rb') as docx_file:
                        result = mammoth.extract_raw_text(docx_file)
                        return result.value
                except Exception as e:
                    print(f"Error reading .doc file with mammoth: {str(e)}")
                    try:
                        # Last resort: try textract
                        return textract.process(file_path).decode('utf-8')
                    except Exception as e:
                        print(f"Error reading .doc file with textract: {str(e)}")
                        return None
    except Exception as e:
        print(f"Error extracting text from Word document: {str(e)}")
        return None

def read_users():
    """Reads user data from the JSON file."""
    if not os.path.exists(USERS_FILE):
        return []
    try:
        with open(USERS_FILE, 'r') as f:
            # Handle empty file case
            content = f.read()
            if not content:
                return []
            # Move the file pointer to the beginning after reading content
            f.seek(0)
            return json.load(f)
    except json.JSONDecodeError:
        # Handle cases where the file is not valid JSON or empty
        return []
    except Exception as e:
        print(f"Error reading users file: {e}")
        return []

def write_users(users):
    """Writes user data to the JSON file."""
    try:
        with open(USERS_FILE, 'w') as f:
            json.dump(users, f, indent=4)
    except Exception as e:
        print(f"Error writing users file: {e}")

# Add this with the other Pydantic models at the top of the file
class JobDescriptionData(BaseModel):
    title: str
    description: str

# Add this with the other Pydantic models at the top of the file
class ResumeAnalysisRequest(BaseModel):
    jobDescription: str
    resumes: List[Dict[str, str]]

# Configure local Mistral endpoint
MISTRAL_ENDPOINT = "http://localhost:11434/api/generate"  # Default Ollama endpoint

def call_mistral(prompt: str) -> str:
    """Call local Mistral model through Ollama."""
    try:
        print("   [Mistral] Sending request to local model...")
        response = requests.post(
            MISTRAL_ENDPOINT,
            json={
                "model": "mistral",
                "prompt": prompt,
                "stream": False
            }
        )
        response.raise_for_status()
        result = response.json()
        print("   [Mistral] Received response from local model")
        return result.get('response', '')
    except Exception as e:
        print(f"   [Mistral] Error calling local model: {str(e)}")
        return ""

@app.post("/api/signup")
async def signup_user_endpoint(user_data: SignupRequest):
    """FastAPI endpoint to handle user signup."""
    users = read_users()

    # Check if email already exists
    for user in users:
        if user.get('email') == user_data.email:
            raise HTTPException(status_code=400, detail="Email already exists.")

    # Validate password match
    if user_data.password != user_data.re_password:
         raise HTTPException(status_code=400, detail="Passwords do not match.")
         
    # Create new user dictionary
    new_user = {
        'role': user_data.role,
        'name': user_data.name,
        'email': user_data.email,
        'password': user_data.password, # Remember to hash this in a real app!
        'phone': user_data.phone,
        'created_at': str(datetime.datetime.now()),
    }

    # Add role-specific fields
    if user_data.role == 'seeker':
        # Check for required seeker fields, allowing None for optional ones defined in model
        if not all([
            user_data.job_title is not None,
            user_data.experience is not None,
            user_data.skills is not None,
            user_data.location is not None,
            user_data.resume_url is not None
        ]):
             raise HTTPException(status_code=400, detail="Missing seeker-specific fields.")
        new_user.update({
            'job_title': user_data.job_title,
            'experience': user_data.experience,
            'skills': user_data.skills,
            'location': user_data.location,
            'resume_url': user_data.resume_url
        })
    elif user_data.role == 'employer':
        # Check for required employer fields, allowing None for optional ones defined in model
        if not all([
            user_data.industry is not None,
            user_data.company_size is not None,
            user_data.company_website is not None,
            user_data.company_location is not None,
            user_data.company_description is not None,
            user_data.contact_person is not None
        ]):
             raise HTTPException(status_code=400, detail="Missing employer-specific fields.")

        new_user.update({
            'industry': user_data.industry,
            'company_size': user_data.company_size,
            'company_website': user_data.company_website,
            'company_location': user_data.company_location,
            'company_description': user_data.company_description,
            'contact_person': user_data.contact_person
        })
    else:
        raise HTTPException(status_code=400, detail="Invalid role specified.")

    users.append(new_user)
    write_users(users)

    return {"message": f"{user_data.role.capitalize()} {user_data.name} signed up successfully!"}

@app.post("/api/login")
async def login_user_endpoint(login_data: LoginRequest):
    """FastAPI endpoint to handle user login."""
    users = read_users()

    # Find the user by email
    user = next((user for user in users if user.get('email') == login_data.email), None)

    # Check if user exists and password matches (plain text comparison for now - hash in real app!)
    if user is None or user.get('password') != login_data.password:
        raise HTTPException(status_code=401, detail="Invalid email or password.")

    # In a real application, you would generate a JWT or session token here
    # For now, just return a success message and the user's role and name
    return {"message": f"Login successful for user {user.get('name')}!", "user": {"name": user.get('name'), "role": user.get('role')}}

# Resume Screening Endpoints

@app.get("/api/job-descriptions")
async def get_job_descriptions():
    try:
        job_descriptions = []
        for filename in os.listdir(JOB_DESCRIPTION_FOLDER):
            if filename.endswith('.txt'):
                filepath = os.path.join(JOB_DESCRIPTION_FOLDER, filename)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                        job_descriptions.append({
                            'name': filename,
                            'title': os.path.splitext(filename)[0],
                            'content': content
                        })
                except Exception as e:
                    print(f"Error reading file {filename}: {str(e)}")
        return job_descriptions
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get job descriptions: {str(e)}")

@app.post("/api/save-job-description")
async def save_job_description(data: JobDescriptionData):
    try:
        if not data.title or not data.description:
            raise HTTPException(status_code=400, detail="Missing required fields")
        
        safe_title = "".join(c for c in data.title if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"{safe_title}.txt"
        filepath = os.path.join(JOB_DESCRIPTION_FOLDER, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(data.description)
        
        return {"message": "Job description saved successfully", "filename": filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save job description: {str(e)}")

@app.get("/api/resumes")
async def list_resumes(path: str = ""):
    """List all resumes in the Original_Resumes directory."""
    try:
        full_path = os.path.join(UPLOAD_FOLDER, path)
        print(f"Listing resumes in path: {full_path}")
        print(f"Path exists: {os.path.exists(full_path)}")
        print(f"Current working directory: {os.getcwd()}")
        print(f"UPLOAD_FOLDER: {UPLOAD_FOLDER}")
        
        if not os.path.exists(full_path):
            print(f"Path does not exist: {full_path}")
            return []
        
        items = []
        for item in os.listdir(full_path):
            item_path = os.path.join(full_path, item)
            # Use forward slashes for consistency
            relative_path = os.path.join(path, item).replace('\\', '/') if path else item
            print(f"Found item: {item_path}")
            print(f"Relative path: {relative_path}")
            
            if os.path.isdir(item_path):
                # Add the folder with trailing slash
                items.append(f"{relative_path}/")
                # Recursively get items from subdirectory
                sub_items = await list_resumes(relative_path)
                items.extend(sub_items)
            elif allowed_file(item):
                items.append(relative_path)
        
        print(f"Returning items: {items}")
        return items
    except Exception as e:
        print(f"Error in list_resumes: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        if not allowed_file(file.filename):
            raise HTTPException(status_code=400, detail="File type not allowed")
        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        with open(filepath, 'wb') as f:
            f.write(await file.read())
        return {"message": "File uploaded successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to upload file: {str(e)}")

@app.post("/api/analyze")
async def analyze_resume_endpoint(data: dict):
    try:
        print(f"Received analyze request with data: {data}")
        
        if not data or 'filename' not in data:
            print("Missing filename in request data")
            raise HTTPException(status_code=400, detail="Missing filename")
            
        filename = data['filename']
        # Check if the file is in the PIXEL Resume directory
        pixel_resume_path = os.path.join(UPLOAD_FOLDER, 'PIXEL Resume', filename)
        if os.path.exists(pixel_resume_path):
            filepath = pixel_resume_path
        else:
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        
        # Debug logging
        print(f"Analyzing file: {filename}")
        print(f"Full filepath: {filepath}")
        print(f"File exists: {os.path.exists(filepath)}")
        print(f"Current working directory: {os.getcwd()}")
        print(f"UPLOAD_FOLDER: {UPLOAD_FOLDER}")
        
        if not os.path.exists(filepath):
            print(f"File not found: {filepath}")
            raise HTTPException(status_code=404, detail=f"File not found: {filepath}")
            
        # Use extract_text_from_word for Word documents
        if filename.lower().endswith(('.doc', '.docx')):
            print("Processing Word document")
            text = extract_text_from_word(filepath)
        else:
            print("Processing non-Word document")
            # For other file types, try to read directly
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                print("UTF-8 encoding failed, trying latin-1")
                # If UTF-8 fails, try with a different encoding
                with open(filepath, 'r', encoding='latin-1') as f:
                    text = f.read()
                    
        if not text:
            print("Failed to extract text from file")
            raise HTTPException(status_code=400, detail="Failed to extract text from file")
            
        # Save parsed resume text
        try:
            parsed_filename = f"parsed_{filename}.txt"
            parsed_filepath = os.path.join(PARSED_RESUMES_FOLDER, parsed_filename)
            with open(parsed_filepath, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Saved parsed resume to: {parsed_filepath}")
        except Exception as e:
            print(f"Warning: Failed to save parsed resume: {str(e)}")
            
        print("Starting resume analysis")
        analysis = analyze_resume(text)
        analysis['parsed_file'] = parsed_filename  # Add the parsed file path to the result
        print("Analysis completed successfully")
        return analysis
    except Exception as e:
        print(f"Error in analyze_resume_endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze resume: {str(e)}")

@app.get("/api/check-ollama")
async def check_ollama():
    try:
        result = test_ollama_connection()
        return {"status": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to check Ollama: {str(e)}")

@app.delete("/api/delete")
async def delete_item(data: dict):
    """Delete a file or folder from the Original_Resumes directory."""
    try:
        path = data.get('path', '')
        is_folder = data.get('isFolder', False)
        
        if not path:
            raise HTTPException(status_code=400, detail="Path is required")

        full_path = os.path.join(UPLOAD_FOLDER, path)
        
        # Check if the path is within the Original_Resumes directory
        if not os.path.abspath(full_path).startswith(os.path.abspath(UPLOAD_FOLDER)):
            raise HTTPException(status_code=403, detail="Access denied: Cannot delete items outside Original_Resumes directory")

        if not os.path.exists(full_path):
            raise HTTPException(status_code=404, detail="Item not found")

        if is_folder:
            shutil.rmtree(full_path)
        else:
            os.remove(full_path)

        return {"message": "Item deleted successfully"}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/move")
async def move_item(data: dict):
    try:
        if not data or 'source' not in data or 'destination' not in data:
            raise HTTPException(status_code=400, detail="Missing source or destination")
        source = os.path.join(BASE_DIR, data['source'])
        destination = os.path.join(BASE_DIR, data['destination'])
        if os.path.exists(source):
            shutil.move(source, destination)
            return {"message": "Item moved successfully"}
        else:
            raise HTTPException(status_code=404, detail="Source item not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to move item: {str(e)}")

@app.post("/api/results")
async def get_results(data: dict):
    """Get analysis results for a resume."""
    try:
        if not data or 'filename' not in data:
            raise HTTPException(status_code=400, detail="Missing filename")
        
        filename = data['filename']
        # Check if the file is in the PIXEL Resume directory
        pixel_resume_path = os.path.join(UPLOAD_FOLDER, 'PIXEL Resume', filename)
        if os.path.exists(pixel_resume_path):
            filepath = pixel_resume_path
        else:
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail=f"File not found: {filepath}")
        
        # Use extract_text_from_word for Word documents
        if filename.lower().endswith(('.doc', '.docx')):
            text = extract_text_from_word(filepath)
        else:
            # For other file types, try to read directly
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # If UTF-8 fails, try with a different encoding
                with open(filepath, 'r', encoding='latin-1') as f:
                    text = f.read()
        
        if not text:
            raise HTTPException(status_code=400, detail="Failed to extract text from file")
        
        analysis = analyze_resume(text)
        # Update the filename in the analysis result
        analysis['filename'] = filename
        return analysis
    except Exception as e:
        print(f"Error in /api/results: {str(e)}")  # Add logging
        raise HTTPException(status_code=500, detail=f"Failed to get results: {str(e)}")

@app.get("/api/resume/{filename}")
async def get_resume_details(filename: str):
    """Get detailed information about a specific resume."""
    try:
        # URL decode the filename
        from urllib.parse import unquote
        filename = unquote(filename)
        
        # Check if the file is in the PIXEL Resume directory
        pixel_resume_path = os.path.join(UPLOAD_FOLDER, 'PIXEL Resume', filename)
        if os.path.exists(pixel_resume_path):
            filepath = pixel_resume_path
        else:
        filepath = os.path.join(UPLOAD_FOLDER, filename)
            
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail=f"File not found: {filename}")
        
        # Use the appropriate text extraction method based on file type
        if filename.lower().endswith(('.doc', '.docx')):
            text = extract_text_from_word(filepath)
        else:
            # For other file types, try to read directly
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # If UTF-8 fails, try with a different encoding
                with open(filepath, 'r', encoding='latin-1') as f:
                    text = f.read()
        
        if not text:
            raise HTTPException(status_code=400, detail="Failed to extract text from file")
        
        # Save parsed resume text
        try:
            parsed_filename = f"parsed_{filename}.txt"
            parsed_filepath = os.path.join(PARSED_RESUMES_FOLDER, parsed_filename)
            with open(parsed_filepath, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Saved parsed resume to: {parsed_filepath}")
        except Exception as e:
            print(f"Warning: Failed to save parsed resume: {str(e)}")
        
        # Get file metadata
        stats = os.stat(filepath)
        file_info = {
            'filename': filename,
            'size': stats.st_size,
            'created': datetime.datetime.fromtimestamp(stats.st_ctime).isoformat(),
            'modified': datetime.datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'content': text[:1000] + '...' if len(text) > 1000 else text,  # Preview of content
            'parsed_file': parsed_filename  # Add the parsed file path to the result
        }
        
        return file_info
    except Exception as e:
        print(f"Error in get_resume_details: {str(e)}")  # Add logging
        raise HTTPException(status_code=500, detail=f"Failed to get resume details: {str(e)}")

@app.get("/api/resume/file/{filename}")
async def get_resume_file(filename: str):
    """Download a specific resume file."""
    try:
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="File not found")
        
        return FileResponse(
            filepath,
            media_type='application/octet-stream',
            filename=filename
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get resume file: {str(e)}")

@app.post("/api/create-folder")
async def create_folder(data: dict):
    """Create a new folder in the Original_Resumes directory."""
    try:
        folder_path = data.get('path', '')
        if not folder_path:
            raise HTTPException(status_code=400, detail="Folder path is required")

        full_path = os.path.join(UPLOAD_FOLDER, folder_path)
        
        # Check if the path is within the Original_Resumes directory
        if not os.path.abspath(full_path).startswith(os.path.abspath(UPLOAD_FOLDER)):
            raise HTTPException(status_code=403, detail="Access denied: Cannot create folder outside Original_Resumes directory")

        if os.path.exists(full_path):
            raise HTTPException(status_code=400, detail="Folder already exists")

        os.makedirs(full_path)
        return {"message": "Folder created successfully"}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/analyze-resumes")
async def analyze_resumes(request: ResumeAnalysisRequest):
    try:
        print("\n=== Starting Resume Analysis ===")
        print(f"Number of resumes to analyze: {len(request.resumes)}")
        print(f"Job description length: {len(request.jobDescription)} characters")
        print("\nJob Description Preview:")
        print("-" * 50)
        print(request.jobDescription[:500] + "..." if len(request.jobDescription) > 500 else request.jobDescription)
        print("-" * 50)
        
        results = []
        for resume in request.resumes:
            print(f"\n--- Analyzing Resume: {resume.get('name', 'Unknown')} ---")
            
            # Get resume content directly
            resume_text = resume.get('content', '')
            print(f"   Resume content length: {len(resume_text)} characters")
            
            # Save parsed resume text
            try:
                parsed_filename = f"parsed_{resume.get('name', 'unknown')}.txt"
                parsed_filepath = os.path.join(PARSED_RESUMES_FOLDER, parsed_filename)
                with open(parsed_filepath, 'w', encoding='utf-8') as f:
                    f.write(resume_text)
                print(f"   Saved parsed resume to: {parsed_filepath}")
            except Exception as e:
                print(f"   Warning: Failed to save parsed resume: {str(e)}")
            
            # Analyze with Mistral
            print("\n2. Starting Mistral analysis...")
            print("   Using job description for analysis:")
            print("   " + "-" * 40)
            print("   " + request.jobDescription[:200] + "..." if len(request.jobDescription) > 200 else "   " + request.jobDescription)
            print("   " + "-" * 40)
            
            print("   - Analyzing skills match...")
            skills_analysis = analyze_skills_match(resume_text, request.jobDescription)
            print(f"   - Found {len(skills_analysis.get('matching_skills', []))} matching skills")
            print(f"   - Found {len(skills_analysis.get('additional_skills', []))} additional valuable skills")
            
            print("   - Analyzing experience match...")
            experience_analysis = analyze_experience_match(resume_text, request.jobDescription)
            print(f"   - Found {len(experience_analysis.get('relevant_experience', []))} relevant experiences")
            print(f"   - Found {len(experience_analysis.get('additional_experience', []))} additional valuable experiences")
            
            print("   - Analyzing education match...")
            education_analysis = analyze_education_match(resume_text, request.jobDescription)
            print(f"   - Education match score: {education_analysis.get('match_score', 0)}%")
            
            # Calculate overall score
            print("3. Calculating overall score...")
            overall_score = calculate_overall_score(skills_analysis, experience_analysis, education_analysis)
            print(f"   Overall score: {overall_score}%")
            
            result = {
                "id": resume.get('id', ''),
                "name": resume.get('name', 'Unknown'),
                "score": overall_score,
                "skills": skills_analysis,
                "experience": experience_analysis,
                "education": education_analysis,
                "raw_text": resume_text,
                "parsed_file": parsed_filename
            }
            results.append(result)
            print(f"--- Completed analysis for {resume.get('name', 'Unknown')} ---\n")
        
        print("=== Resume Analysis Complete ===\n")
        return results
        
    except Exception as e:
        print(f"\n!!! Error in resume analysis: {str(e)}")
        print(f"Error type: {type(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

def extract_text_with_mistral(content: str) -> str:
    """Extract and clean text from resume content using Mistral."""
    try:
        print("   [Mistral] Starting text extraction...")
        if not content:
            print("   [Mistral] Warning: Empty content received")
            return ""
            
        # If content is already text, return it
        if isinstance(content, str) and not content.startswith('%PDF'):
            print("   [Mistral] Content is already text, cleaning it...")
            return clean_resume_text(content)
            
        # Prepare prompt for Mistral
        prompt = f"""Extract and clean the text from this resume content. 
        Remove any formatting, special characters, or unnecessary whitespace.
        Return ONLY the cleaned text, without any additional comments or formatting.

        Resume Content:
        {content}

        Cleaned Text:"""

        # Get response from Mistral
        response = call_mistral(prompt)
        cleaned_text = response.strip()
        
        print(f"   [Mistral] Successfully extracted {len(cleaned_text)} characters")
        return cleaned_text
    except Exception as e:
        print(f"   [Mistral] Error extracting text: {str(e)}")
        return content if isinstance(content, str) else ""

def analyze_skills_match(resume_text: str, job_description: str) -> Dict:
    """Analyze skills match between resume and job description using local Mistral."""
    print("   [Mistral] Starting skills analysis...")
    print("   [Mistral] Using job description for skills matching:")
    print("   [Mistral] " + "-" * 40)
    print("   [Mistral] " + job_description[:200] + "..." if len(job_description) > 200 else "   [Mistral] " + job_description)
    print("   [Mistral] " + "-" * 40)
    try:
        # Prepare the prompt for Mistral
        prompt = f"""First, analyze the resume to extract all skills mentioned. Then, compare with the job description.
        Return a JSON object with the following structure:
        {{
            "match_score": 0-100,
            "matching_skills": [
                {{
                    "name": "skill name (exactly as mentioned in resume)",
                    "level": "beginner/intermediate/expert (based on resume context)",
                    "years": "years of experience (from resume)",
                    "context": "how this skill is mentioned in the resume"
                }},
                ...
            ],
            "additional_skills": [
                {{
                    "name": "skill name (exactly as mentioned in resume)",
                    "level": "beginner/intermediate/expert (based on resume context)",
                    "years": "years of experience (from resume)",
                    "context": "how this skill is mentioned in the resume"
                }},
                ...
            ],
            "missing_skills": [
                {{
                    "name": "skill name (from job description)",
                    "importance": "high/medium/low",
                    "reason": "why this skill is important for the role",
                    "is_required": true/false
                }},
                ...
            ],
            "skills_summary": "summary of all skills found in the resume and missing requirements"
        }}

        Resume:
        {resume_text}

        Job Description:
        {job_description}

        Important: 
        1. First, extract ALL skills mentioned in the resume
        2. Only include skills that are explicitly mentioned in the resume
        3. Keep the original wording from the resume
        4. Then, identify skills required in the job description that are not in the resume
        5. Mark each missing skill as required or preferred based on job description
        6. Focus on what the candidate actually has, but also note what's missing

        Analysis:"""

        # Get response from local Mistral
        response = call_mistral(prompt)
        
        # Parse the JSON response
        try:
            result = json.loads(response)
            print(f"   [Mistral] Extracted {len(result.get('matching_skills', []))} matching skills")
            return result
        except json.JSONDecodeError:
            print("   [Mistral] Error parsing skills analysis response")
            return {
                "match_score": 0,
                "matching_skills": [],
                "additional_skills": [],
                "missing_skills": [],
                "skills_summary": "Unable to analyze skills"
            }
    except Exception as e:
        print(f"   [Mistral] Error in skills analysis: {str(e)}")
        return {
            "match_score": 0,
            "matching_skills": [],
            "additional_skills": [],
            "missing_skills": [],
            "skills_summary": "Error analyzing skills"
        }

def analyze_experience_match(resume_text: str, job_description: str) -> Dict:
    """Analyze experience match between resume and job description using local Mistral."""
    print("   [Mistral] Starting experience analysis...")
    print("   [Mistral] Using job description for experience matching:")
    print("   [Mistral] " + "-" * 40)
    print("   [Mistral] " + job_description[:200] + "..." if len(job_description) > 200 else "   [Mistral] " + job_description)
    print("   [Mistral] " + "-" * 40)
    try:
        # Prepare the prompt for Mistral
        prompt = f"""First, analyze the resume to extract all work experience. Then, compare with the job description.
        Return a JSON object with the following structure:
        {{
            "match_score": 0-100,
            "relevant_experience": [
                {{
                    "role": "job title (exactly as mentioned in resume)",
                    "company": "company name (exactly as mentioned in resume)",
                    "duration": "time period (exactly as mentioned in resume)",
                    "responsibilities": ["responsibility1", "responsibility2", ...] (from resume),
                    "achievements": ["achievement1", "achievement2", ...] (from resume)
                }},
                ...
            ],
            "additional_experience": [
                {{
                    "role": "job title (exactly as mentioned in resume)",
                    "company": "company name (exactly as mentioned in resume)",
                    "duration": "time period (exactly as mentioned in resume)",
                    "responsibilities": ["responsibility1", "responsibility2", ...] (from resume),
                    "achievements": ["achievement1", "achievement2", ...] (from resume)
                }},
                ...
            ],
            "missing_experience": [
                {{
                    "requirement": "experience requirement (from job description)",
                    "importance": "high/medium/low",
                    "reason": "why this experience is important",
                    "is_required": true/false
                }},
                ...
            ],
            "experience_summary": "summary of all work experience found in the resume and missing requirements"
        }}

        Resume:
        {resume_text}

        Job Description:
        {job_description}

        Important: 
        1. First, extract ALL work experience mentioned in the resume
        2. Only include experience that is explicitly mentioned in the resume
        3. Keep the original wording from the resume
        4. Then, identify experience required in the job description that is not in the resume
        5. Mark each missing experience as required or preferred based on job description
        6. Focus on what the candidate actually has, but also note what's missing

        Analysis:"""

        # Get response from local Mistral
        response = call_mistral(prompt)
        
        # Parse the JSON response
        try:
            result = json.loads(response)
            print(f"   [Mistral] Calculated experience match: {result.get('match_score', 0)}%")
            return result
        except json.JSONDecodeError:
            print("   [Mistral] Error parsing experience analysis response")
            return {
                "match_score": 0,
                "relevant_experience": [],
                "additional_experience": [],
                "missing_experience": [],
                "experience_summary": "Unable to analyze experience"
            }
    except Exception as e:
        print(f"   [Mistral] Error in experience analysis: {str(e)}")
        return {
            "match_score": 0,
            "relevant_experience": [],
            "additional_experience": [],
            "missing_experience": [],
            "experience_summary": "Error analyzing experience"
        }

def analyze_education_match(resume_text: str, job_description: str) -> Dict:
    """Analyze education match between resume and job description using local Mistral."""
    print("   [Mistral] Starting education analysis...")
    print("   [Mistral] Using job description for education matching:")
    print("   [Mistral] " + "-" * 40)
    print("   [Mistral] " + job_description[:200] + "..." if len(job_description) > 200 else "   [Mistral] " + job_description)
    print("   [Mistral] " + "-" * 40)
    try:
        # Prepare the prompt for Mistral
        prompt = f"""First, analyze the resume to extract all education details. Then, compare with the job description.
        Return a JSON object with the following structure:
        {{
            "match_score": 0-100,
            "education_details": [
                {{
                    "degree": "degree name (exactly as mentioned in resume)",
                    "field": "field of study (exactly as mentioned in resume)",
                    "institution": "school name (exactly as mentioned in resume)",
                    "graduation_year": "year (exactly as mentioned in resume)",
                    "gpa": "GPA if available (exactly as mentioned in resume)",
                    "achievements": ["achievement1", "achievement2", ...] (from resume)
                }},
                ...
            ],
            "missing_education": [
                {{
                    "requirement": "education requirement (from job description)",
                    "importance": "high/medium/low",
                    "reason": "why this education is important",
                    "is_required": true/false
                }},
                ...
            ],
            "education_summary": "summary of all education found in the resume and missing requirements"
        }}

        Resume:
        {resume_text}

        Job Description:
        {job_description}

        Important: 
        1. First, extract ALL education details mentioned in the resume
        2. Only include education that is explicitly mentioned in the resume
        3. Keep the original wording from the resume
        4. Then, identify education required in the job description that is not in the resume
        5. Mark each missing education as required or preferred based on job description
        6. Focus on what the candidate actually has, but also note what's missing

        Analysis:"""

        # Get response from local Mistral
        response = call_mistral(prompt)
        
        # Parse the JSON response
        try:
            result = json.loads(response)
            print(f"   [Mistral] Calculated education match: {result.get('match_score', 0)}%")
            return result
        except json.JSONDecodeError:
            print("   [Mistral] Error parsing education analysis response")
            return {
                "match_score": 0,
                "education_details": [],
                "missing_education": [],
                "education_summary": "Unable to analyze education"
            }
    except Exception as e:
        print(f"   [Mistral] Error in education analysis: {str(e)}")
        return {
            "match_score": 0,
            "education_details": [],
            "missing_education": [],
            "education_summary": "Error analyzing education"
        }

def calculate_overall_score(skills_analysis: Dict, experience_analysis: Dict, education_analysis: Dict) -> int:
    """Calculate overall match score with bonus points for additional valuable skills/experience."""
    try:
        # Base weights
        skills_weight = 0.4  # 40% weight for skills
        experience_weight = 0.4  # 40% weight for experience
        education_weight = 0.2  # 20% weight for education
        
        # Calculate base scores - ensure they are integers
        skills_score = int(skills_analysis.get('match_score', 0))
        experience_score = int(experience_analysis.get('match_score', 0))
        education_score = int(education_analysis.get('match_score', 0))
        
        # Calculate bonus points for additional valuable skills/experience
        skills_bonus = min(len(skills_analysis.get('additional_skills', [])) * 5, 20)  # Up to 20% bonus
        experience_bonus = min(len(experience_analysis.get('additional_experience', [])) * 5, 20)  # Up to 20% bonus
        
        # Calculate weighted average with bonuses
        overall_score = (
            (skills_score + skills_bonus) * skills_weight +
            (experience_score + experience_bonus) * experience_weight +
            education_score * education_weight
        )
        
        return min(round(overall_score), 100)  # Cap at 100%
    except Exception as e:
        print(f"Error calculating overall score: {str(e)}")
        return 0

# You would typically run this with uvicorn:
# python -m uvicorn backend.main:app --reload 