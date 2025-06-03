import os
import re
import json
import requests
import PyPDF2
import docx
import textract
import pandas as pd
import io
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import mammoth
import chardet
from typing import Dict, Any, Set, List, Tuple, Optional
from docx import Document
import datetime

# Configure parsed resumes folder (temporary)
PARSED_RESUMES_FOLDER = 'parsed_resumes'
if not os.path.exists(PARSED_RESUMES_FOLDER):
    os.makedirs(PARSED_RESUMES_FOLDER)

# Configure Tesseract path for Windows
def configure_tesseract():
    """Configure Tesseract OCR path and settings."""
    try:
        # Common Tesseract installation paths on Windows
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        ]
        
        # Try to find Tesseract installation
        tesseract_path = None
        for path in possible_paths:
            if os.path.exists(path):
                tesseract_path = path
                break
        
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            print(f"Tesseract configured successfully at: {tesseract_path}")
            return True
        else:
            print("Tesseract not found in common installation paths.")
            print("Please set the Tesseract path manually in the code.")
            return False
            
    except Exception as e:
        print(f"Error configuring Tesseract: {str(e)}")
        return False

# Call configuration at startup
configure_tesseract()

def parse_date(date_str: str) -> Optional[Tuple[int, int]]:
    """Parse date string to (year, month) tuple"""
    try:
        # Month to number mapping
        month_map = {
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12,
            'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
            'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
        }
        
        # Handle month-year format
        for month_name, month_num in month_map.items():
            if month_name in date_str.lower():
                year_match = re.search(r'(?:19|20)\d{2}', date_str)
                if year_match:
                    year = int(year_match.group())
                    return (year, month_num)
        
        # Handle numeric date formats
        if '/' in date_str or '-' in date_str:
            parts = re.split(r'[/-]', date_str)
            if len(parts) == 3:
                try:
                    # MM/DD/YYYY
                    month, day, year = map(int, parts)
                    if len(str(year)) == 2:
                        year += 2000
                    return (year, month)
                except ValueError:
                    try:
                        # DD/MM/YYYY
                        day, month, year = map(int, parts)
                        if len(str(year)) == 2:
                            year += 2000
                        return (year, month)
                    except ValueError:
                        return None
        
        # Handle year-only format
        year_match = re.search(r'(?:19|20)\d{2}', date_str)
        if year_match:
            year = int(year_match.group())
            return (year, 1)  # Default to January if only year is provided
            
        return None
    except Exception as e:
        print(f"Error parsing date '{date_str}': {str(e)}")
        return None

def calculate_months_between(start_date: Tuple[int, int], end_date: Tuple[int, int]) -> int:
    """Calculate months between two dates"""
    if not start_date or not end_date:
        return 0
    start_year, start_month = start_date
    end_year, end_month = end_date
    
    return (end_year - start_year) * 12 + (end_month - start_month)

def extract_text_from_resume(file_path: str) -> Optional[str]:
    """Extract text from resume file based on its extension."""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            return extract_text_from_word(file_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(file_path)
        elif file_extension == '.rtf':
            return extract_text_from_rtf(file_path)
        elif file_extension == '.odt':
            return extract_text_from_odt(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png']:
            return extract_text_from_image(file_path)
        else:
            print(f"Unsupported file type: {file_extension}")
            return None
            
    except Exception as e:
        print(f"Error extracting text from resume: {str(e)}")
        return None

def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF with OCR fallback."""
    try:
        text = ""
        with open(filepath, 'rb') as file:
            # Try PyPDF2 first
            try:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if not page_text.strip():  # If no text extracted, try OCR
                        try:
                            images = convert_from_path(filepath, first_page=page.page_number, last_page=page.page_number)
                            if images:
                                page_text = pytesseract.image_to_string(images[0])
                        except Exception as e:
                            print(f"OCR fallback failed for page {page.page_number}: {str(e)}")
                    text += page_text + "\n"
            except Exception as e:
                print(f"PyPDF2 extraction failed: {str(e)}")
                # Fallback to full document OCR
                try:
                    images = convert_from_path(filepath)
                    for image in images:
                        text += pytesseract.image_to_string(image) + "\n"
                except Exception as e:
                    print(f"PDF image extraction failed: {str(e)}")
        
        return text.strip()
    except Exception as e:
        print(f"Error in PDF extraction: {str(e)}")
        return ""

def extract_text_from_word(filepath: str) -> str:
    """Extract text from Word documents."""
    try:
        # For .docx files
        if filepath.endswith('.docx'):
            doc = Document(filepath)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        # For .doc files
        elif filepath.endswith('.doc'):
            try:
                # Try using python-docx directly
                doc = Document(filepath)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
            except Exception as e:
                print(f"Error reading .doc file: {str(e)}")
                # If direct reading fails, try using mammoth
                try:
                    with open(filepath, 'rb') as docx_file:
                        result = mammoth.extract_raw_text(docx_file)
                        return result.value
                except Exception as e:
                    print(f"Error using mammoth for .doc file: {str(e)}")
                    return ""
    except Exception as e:
        print(f"Error in Word extraction: {str(e)}")
        return ""

def extract_text_from_txt(filepath: str) -> str:
    """Extract text from text files with encoding detection."""
    try:
        # Detect file encoding
        with open(filepath, 'rb') as file:
            raw_data = file.read()
            detected = chardet.detect(raw_data)
            encoding = detected['encoding']
        
        # Read file with detected encoding
        with open(filepath, 'r', encoding=encoding) as file:
            return file.read()
    except Exception as e:
        print(f"Error in text file extraction: {str(e)}")
        return ""

def extract_text_from_rtf(filepath: str) -> str:
    """Extract text from RTF files."""
    try:
        return textract.process(filepath).decode('utf-8')
    except Exception as e:
        print(f"Error in RTF extraction: {str(e)}")
        return ""

def extract_text_from_odt(filepath: str) -> str:
    """Extract text from ODT files."""
    try:
        return textract.process(filepath).decode('utf-8')
    except Exception as e:
        print(f"Error in ODT extraction: {str(e)}")
        return ""

def extract_text_from_image(filepath: str) -> str:
    """Extract text from image files using OCR."""
    try:
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        print(f"Error in image extraction: {str(e)}")
        return ""

def clean_resume_text(text: str) -> str:
    """Clean the extracted resume text."""
    if not text:
        return ""
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep important ones
    text = re.sub(r'[^\w\s.,;:()@-]', '', text)
    
    # Remove multiple newlines
    text = re.sub(r'\n\s*\n', '\n', text)
    
    # Remove multiple spaces
    text = re.sub(r'\s{2,}', ' ', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information from resume text."""
    contact = {}
    
    # Extract email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact['email'] = email_match.group()
    
    # Extract phone
    phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
    phone_match = re.search(phone_pattern, text)
    if phone_match:
        contact['phone'] = phone_match.group()
    
    # Extract location
    location_pattern = r'\b[A-Z][a-z]+(?:[\s-][A-Z][a-z]+)*,\s*[A-Z]{2}\b'
    location_match = re.search(location_pattern, text)
    if location_match:
        contact['location'] = location_match.group()
    
    return contact

def extract_skills(text: str) -> Dict[str, Set[str]]:
    """Extract skills from resume text."""
    skills = {
        'technical_skills': set(),
        'tools': set(),
        'soft_skills': set(),
        'languages': set()
    }
    
    # Common technical skills
    technical_skills = [
        'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift',
        'kotlin', 'rust', 'go', 'typescript', 'html', 'css', 'sql', 'nosql',
        'mongodb', 'postgresql', 'mysql', 'oracle', 'aws', 'azure', 'gcp',
        'docker', 'kubernetes', 'react', 'angular', 'vue', 'node.js', 'express',
        'django', 'flask', 'spring', 'asp.net', 'machine learning', 'ai',
        'data science', 'big data', 'hadoop', 'spark', 'tensorflow', 'pytorch'
    ]
    
    # Common tools
    tools = [
        'git', 'github', 'gitlab', 'bitbucket', 'jenkins', 'jira', 'confluence',
        'slack', 'teams', 'zoom', 'figma', 'sketch', 'photoshop', 'illustrator',
        'excel', 'powerpoint', 'word', 'outlook', 'vscode', 'intellij', 'eclipse',
        'postman', 'swagger', 'docker', 'kubernetes', 'terraform', 'ansible'
    ]
    
    # Common soft skills
    soft_skills = [
        'leadership', 'communication', 'teamwork', 'problem solving',
        'critical thinking', 'time management', 'project management',
        'agile', 'scrum', 'collaboration', 'adaptability', 'creativity',
        'analytical', 'strategic thinking', 'decision making'
    ]
    
    # Common languages
    languages = [
        'english', 'spanish', 'french', 'german', 'chinese', 'japanese',
        'korean', 'russian', 'arabic', 'hindi', 'portuguese', 'italian'
    ]
    
    # Convert text to lowercase for case-insensitive matching
    text_lower = text.lower()
    
    # Extract skills
    for skill in technical_skills:
        if skill in text_lower:
            skills['technical_skills'].add(skill)
    
    for tool in tools:
        if tool in text_lower:
            skills['tools'].add(tool)
    
    for skill in soft_skills:
        if skill in text_lower:
            skills['soft_skills'].add(skill)
    
    for language in languages:
        if language in text_lower:
            skills['languages'].add(language)
    
    return skills

def extract_education(text: str) -> List[Dict[str, str]]:
    """Extract education information from resume text."""
    education = []
    
    # Common degree patterns
    degree_patterns = [
        r'\b(?:Bachelor|Master|PhD|Doctorate|B\.S\.|M\.S\.|B\.A\.|M\.A\.|B\.E\.|M\.E\.|B\.Tech|M\.Tech)\b',
        r'\b(?:Associate|Diploma|Certificate)\b'
    ]
    
    # Find all degree mentions
    for pattern in degree_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            # Get some context around the degree
            start = max(0, match.start() - 50)
            end = min(len(text), match.end() + 50)
            context = text[start:end]
            
            # Try to extract institution name
            institution_pattern = r'\b(?:University|College|Institute|School)\s+of\s+[A-Za-z\s]+\b'
            institution_match = re.search(institution_pattern, context)
            
            education.append({
                'degree': match.group(),
                'institution': institution_match.group() if institution_match else 'Not specified',
                'field': 'Not specified',
                'start_date': 'Not specified',
                'end_date': 'Not specified',
                'gpa': 'Not specified'
            })
    
    return education

def analyze_resume(text: str, job_description: str = '') -> Dict[str, Any]:
    """Analyze the resume text using Mistral LLM and compare with job description."""
    try:
        print("\n=== Starting Resume Analysis ===")
        print("Start Time:", datetime.datetime.now().strftime("%H:%M:%S"))
        print(f"Job Description Length: {len(job_description)} characters")
        print(f"Resume Text Length: {len(text)} characters")
        
        # First try to get analysis from Mistral
        print("\nCalling Mistral API for analysis...")
        mistral_response = call_mistral_api(text, job_description)
        
        if not mistral_response:
            print("Mistral API not available, performing basic analysis")
            skills = extract_skills(text)
            print(f"Extracted {len(skills.get('technical_skills', set()))} technical skills")
            return {
                'overall_score': 0,
                'skills': list(skills.get('technical_skills', set())),
                'experience_match': 0,
                'education_match': 0
            }
            
        # Parse the JSON response from Mistral
        try:
            print("\nParsing Mistral response...")
            analysis = json.loads(mistral_response)
            skills = analysis.get('skills', {})
            
            print("\nCalculating match scores...")
            experience_match = calculate_experience_match(analysis.get('experience', []), job_description)
            education_match = calculate_education_match(analysis.get('education', []), job_description)
            
            print(f"Experience Match: {experience_match}%")
            print(f"Education Match: {education_match}%")
            print(f"Technical Skills Found: {len(skills.get('technical_skills', []))}")
            
            result = {
                'overall_score': analysis.get('overall_score', 0),
                'skills': list(skills.get('technical_skills', [])),
                'experience_match': experience_match,
                'education_match': education_match
            }
            
            print("\nAnalysis Results:")
            print(f"Overall Score: {result['overall_score']}%")
            print(f"Skills: {', '.join(result['skills'])}")
            print(f"Experience Match: {result['experience_match']}%")
            print(f"Education Match: {result['education_match']}%")
            
            print("\n=== Resume Analysis Complete ===")
            print("End Time:", datetime.datetime.now().strftime("%H:%M:%S"))
            
            return result
            
        except json.JSONDecodeError as e:
            print(f"Failed to parse Mistral response as JSON: {str(e)}")
            skills = extract_skills(text)
            return {
                'overall_score': 0,
                'skills': list(skills.get('technical_skills', set())),
                'experience_match': 0,
                'education_match': 0
            }
    except Exception as e:
        print(f"Error in resume analysis: {str(e)}")
        print("Stack trace:", e.__traceback__)
        skills = extract_skills(text)
        return {
            'overall_score': 0,
            'skills': list(skills.get('technical_skills', set())),
            'experience_match': 0,
            'education_match': 0
        }

def calculate_experience_match(experience: List[Dict[str, Any]], job_description: str) -> int:
    """Calculate experience match percentage."""
    if not experience or not job_description:
        return 0
    
    # Extract key terms from job description
    job_terms = set(re.findall(r'\b\w+\b', job_description.lower()))
    
    # Count matching terms in experience
    match_count = 0
    total_terms = len(job_terms)
    
    for exp in experience:
        exp_text = ' '.join([
            exp.get('company', ''),
            exp.get('role', ''),
            ' '.join(exp.get('responsibilities', []))
        ]).lower()
        
        for term in job_terms:
            if term in exp_text:
                match_count += 1
    
    return int((match_count / total_terms) * 100) if total_terms > 0 else 0

def calculate_education_match(education: List[Dict[str, Any]], job_description: str) -> int:
    """Calculate education match percentage."""
    if not education or not job_description:
        return 0
    
    # Extract key terms from job description
    job_terms = set(re.findall(r'\b\w+\b', job_description.lower()))
    
    # Count matching terms in education
    match_count = 0
    total_terms = len(job_terms)
    
    for edu in education:
        edu_text = ' '.join([
            edu.get('degree', ''),
            edu.get('institution', ''),
            edu.get('field', '')
        ]).lower()
        
        for term in job_terms:
            if term in edu_text:
                match_count += 1
    
    return int((match_count / total_terms) * 100) if total_terms > 0 else 0

def call_mistral_api(prompt: str, job_description: str = '') -> str:
    """Call the Mistral API for text analysis with job description comparison."""
    try:
        # Configure the API endpoint and headers
        url = "http://localhost:11434/api/generate"
        headers = {
            "Content-Type": "application/json"
        }
        
        # Prepare the prompt for resume analysis
        system_prompt = """You are an expert resume analyzer and job matching specialist. Analyze the given resume text against the job description and extract:
1. Contact information (email, phone, location)
2. Skills match (technical, tools, soft skills, languages) - compare with job requirements
3. Education match (degree, institution, field, dates, GPA) - compare with job requirements
4. Experience match (company, role, duration, responsibilities) - compare with job requirements
5. Overall assessment and score (0-100) based on job requirements match

Format the response as a JSON object with the following structure:
{
    "contact": {"email": "", "phone": "", "location": ""},
    "skills": {
        "technical_skills": [],
        "tools": [],
        "soft_skills": [],
        "languages": []
    },
    "education": [
        {
            "degree": "",
            "institution": "",
            "field": "",
            "start_date": "",
            "end_date": "",
            "gpa": ""
        }
    ],
    "experience": [
        {
            "company": "",
            "role": "",
            "duration": "",
            "responsibilities": []
        }
    ],
    "overall_score": 0,
    "assessment": ""
}

Calculate the overall score based on:
- Skills match (40%): How well the candidate's skills match the job requirements
- Experience match (30%): How relevant the candidate's experience is to the job
- Education match (20%): How well the candidate's education aligns with the job requirements
- Other factors (10%): Location, availability, etc."""

        # Prepare the request payload
        payload = {
            "model": "mistral",
            "prompt": f"{system_prompt}\n\nJob Description:\n{job_description}\n\nResume Text:\n{prompt}",
            "stream": False
        }
        
        # Make the API call
        response = requests.post(url, headers=headers, json=payload)
        
        if response.status_code == 200:
            result = response.json()
            return result.get('response', '')
        else:
            print(f"Error calling Mistral API: {response.status_code}")
            return ""
            
    except Exception as e:
        print(f"Error in Mistral API call: {str(e)}")
        return ""

def test_ollama_connection() -> bool:
    """Test connection to Ollama service."""
    try:
        response = requests.get('http://localhost:11434/api/tags')
        return response.status_code == 200
    except:
        return False

def test_mistral_functionality():
    """Test if Mistral is working properly with a simple prompt."""
    try:
        print("\n=== Testing Mistral Functionality ===")
        print("Start Time:", datetime.datetime.now().strftime("%H:%M:%S"))
        
        # Simple test prompt
        test_prompt = """Please respond with a simple JSON object: {"status": "working", "message": "Mistral is functioning correctly"}"""
        
        print("Sending test request to Mistral...")
        response = requests.post('http://localhost:11434/api/generate',
                               json={
                                   "model": "mistral",
                                   "prompt": test_prompt,
                                   "stream": False
                               })
        
        print("Response received at:", datetime.datetime.now().strftime("%H:%M:%S"))
        
        if response.status_code == 200:
            try:
                result = response.json()
                response_text = result['response']
                print("\nMistral Response:")
                print(response_text)
                
                # Try to parse JSON from response
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                
                if json_start >= 0 and json_end > json_start:
                    json_str = response_text[json_start:json_end]
                    parsed_data = json.loads(json_str)
                    print("\nSuccessfully parsed JSON response")
                    print("Mistral is working correctly!")
                    return True
                else:
                    print("\nError: Could not find JSON in response")
                    return False
                    
            except Exception as e:
                print(f"\nError processing Mistral response: {str(e)}")
                return False
        else:
            print(f"\nError: Mistral API returned status code {response.status_code}")
            print(f"Response: {response.text}")
            return False
            
    except Exception as e:
        print(f"\nError testing Mistral: {str(e)}")
        return False

if __name__ == '__main__':
    # Test Ollama connection and Mistral functionality
    if not test_ollama_connection():
        print("Please fix the Ollama connection issues before proceeding.")
        exit(1)
        
    print("\nTesting Mistral functionality...")
    if test_mistral_functionality():
        print("\nAll tests passed! Mistral is ready to use.")
    else:
        print("\nMistral test failed. Please check the error messages above.")
        exit(1)

# Placeholder for additional functions or logic 